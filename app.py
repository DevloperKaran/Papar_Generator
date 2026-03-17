from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
import os
import io

app = Flask(__name__)

BASE_FOLDER = "question_bank"
questions = {}

def load_questions(folder):

    global questions
    questions = {}
    i = 0

    path = os.path.join(BASE_FOLDER, folder)

    if not os.path.exists(path):
        return {}

    for file in os.listdir(path):

        if file.endswith(".docx"):

            doc = Document(os.path.join(path, file))

            for para in doc.paragraphs:

                text = para.text.strip()

                if text:
                    qid = f"q{i}"
                    questions[qid] = text
                    i += 1

    return questions


@app.route("/")
def index():

    folders = os.listdir(BASE_FOLDER)

    return render_template("index.html", folders=folders)


@app.route("/load/<folder>")
def load_folder(folder):

    qs = load_questions(folder)

    return jsonify(qs)


@app.route("/generate", methods=["POST"])
def generate():

    data = request.json
    ids = data.get("ids", [])

    doc = Document()

    for qid in ids:

        if qid in questions:
            doc.add_paragraph(questions[qid])

    buffer = io.BytesIO()

    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="question_paper.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


if __name__ == "__main__":
    app.run(debug=True)